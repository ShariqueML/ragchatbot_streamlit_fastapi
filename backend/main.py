from fastapi import FastAPI, UploadFile, File, HTTPException
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any, TypedDict,Generic, TypeVar
from abc import ABC
import uuid
import io
import PyPDF2
from docx import Document as dx
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_ollama.chat_models import ChatOllama
from langchain_core.prompts import PromptTemplate
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langgraph.checkpoint.memory import MemorySaver
from langgraph.graph import StateGraph, END
from sqlalchemy import create_engine, Column, String, Integer, DateTime, ForeignKey, Text
from sqlalchemy.dialects.sqlite import JSON as SQLiteJSON
# from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, relationship
import datetime
from enum import Enum as PyEnum
from contextlib import asynccontextmanager
import uvicorn
from sqlalchemy.orm import DeclarativeBase
# Define the FastAPI application
app = FastAPI(
    title="Conversational RAG API",
    description="A backend for a Conversational RAG Chatbot using LangChain and LangGraph.",
    version="1.0.0"
)
T = TypeVar("T")
# --- 1. Database Setup ---
DATABASE_URL = "sqlite:///database_telemetry.db"
engine = create_engine(DATABASE_URL)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

class Base(DeclarativeBase):
    pass

class FeedbackScore(PyEnum):
    POSITIVE = 1
    NEGATIVE = -1

class Telemetry(Base):
    __tablename__ = "telemetry_table"
    transaction_id = Column(String, primary_key=True)
    session_id = Column(String)
    user_question = Column(Text)
    response = Column(Text)
    context = Column(Text)
    model_name = Column(String)
    input_tokens = Column(Integer)
    output_tokens = Column(Integer)
    total_tokens = Column(Integer)
    latency = Column(Integer)
    dtcreatedon = Column(DateTime)
    
    feedback = relationship("Feedback", back_populates="telemetry_entry", uselist=False)

class Feedback(Base):
    __tablename__ = "feedback_table"
    id = Column(Integer, primary_key=True, autoincrement=True)
    telemetry_entry_id = Column(String, ForeignKey("telemetry_table.transaction_id"), nullable=False, unique=True)
    feedback_score = Column(Integer, nullable=False)
    feedback_text = Column(Text, nullable=True)
    user_query = Column(Text, nullable=False)
    llm_response = Column(Text, nullable=False)
    timestamp = Column(DateTime, default=datetime.datetime.now)
    
    telemetry_entry = relationship("Telemetry", back_populates="feedback")

class ConversationHistory(Base):
    __tablename__ = "conversation_history"
    session_id = Column(String, primary_key=True)
    messages = Column(SQLiteJSON, nullable=False)
    last_updated = Column(DateTime, default=datetime.datetime.now)

# --- 2. Initialize LLM and Embeddings ---
my_model_name = "granite3.3:2b"
llm = ChatOllama(model=my_model_name)
embeddings = HuggingFaceEmbeddings(
    model_name="intfloat/e5-base-v2",
    model_kwargs={'device': 'cpu'},
    encode_kwargs={'normalize_embeddings': False}
)

# --- 3. LangGraph State and Workflow ---
class GraphState(TypedDict):
    chat_history: List[Dict[str, Any]]
    retrieved_documents: List[str]
    user_question: str
    session_id: str
    telemetry_id: Optional[str] = None

vectorstore_retriever = None
compiled_app = None
memory = MemorySaver()

# --- 4. LangGraph Nodes ---
def retrieve_documents(state: GraphState):
    global vectorstore_retriever
    user_question = state["user_question"]
    if vectorstore_retriever is None:
        raise ValueError("Knowledge base not loaded. Please upload documents first.")

    retrieved_docs = vectorstore_retriever.invoke(user_question)
    retrieved_docs_content = [doc.page_content for doc in retrieved_docs]
    return {"retrieved_documents": retrieved_docs_content}

def generate_response(state: GraphState):
    global llm
    user_question = state["user_question"]
    retrieved_documents = state["retrieved_documents"]
    
    formatted_chat_history = []
    for msg in state["chat_history"]:
        if msg['role'] == 'user':
            formatted_chat_history.append(HumanMessage(content=msg['content']))
        elif msg['role'] == 'assistant':
            formatted_chat_history.append(AIMessage(content=msg['content']))
    
    if not retrieved_documents:
        response_content = "I couldn't find any relevant information in the uploaded documents for your question. Can you please rephrase or provide more context?"
        response_obj = AIMessage(content=response_content)
    else:
        context = "\n\n".join(retrieved_documents)
        template = """
            You are a helpful AI assistant. Answer the user's question based on the provided context {context} and the conversation history {chat_history}.
            If the answer is not in the context, state that you don't have enough information.
            Do not make up answers. Only use the given context and chat_history.
            Remove unwanted words like 'Response:' or 'Answer:' from answers.
            \n\nHere is the Question:\n{user_question}
        """
        rag_prompt = PromptTemplate(
            input_variables=["context", "chat_history", "user_question"],
            template=template
        )
        rag_chain = rag_prompt | llm
        response_obj = rag_chain.invoke({
            "context": [SystemMessage(content=context)],
            "chat_history": formatted_chat_history,
            "user_question": [HumanMessage(content=user_question)]
        })
    
    telemetry_data = response_obj.model_dump()
    input_tokens = telemetry_data.get('usage_metadata', {}).get('input_tokens', 0)
    output_tokens = telemetry_data.get('usage_metadata', {}).get('output_tokens', 0)
    total_tokens = telemetry_data.get('usage_metadata', {}).get('total_tokens', 0)
    model_name = telemetry_data.get('response_metadata', {}).get('model', 'unknown')
    total_duration = telemetry_data.get('response_metadata', {}).get('total_duration', 0)
    
    db = SessionLocal()
    transaction_id = str(uuid.uuid4())
    try:
        telemetry_record = Telemetry(
            transaction_id=transaction_id,
            session_id=state.get("session_id"),
            user_question=user_question,
            response=response_obj.content,
            context="\n\n".join(retrieved_documents) if retrieved_documents else "No documents retrieved",
            model_name=model_name,
            input_tokens=input_tokens,
            output_tokens=output_tokens,
            total_tokens=total_tokens,
            latency=total_duration,
            dtcreatedon=datetime.datetime.now()
        )
        db.add(telemetry_record)
        
        new_messages = state["chat_history"] + [
            {"role": "user", "content": user_question}, 
            {"role": "assistant", "content": response_obj.content, "telemetry_id": transaction_id}
        ]
        
        # --- FIX: Refactored Database Save Logic ---
        print(f"Saving conversation for session_id: {state.get('session_id')}")
        conversation_entry = db.query(ConversationHistory).filter_by(session_id=state.get("session_id")).first()
        if conversation_entry:
            print(f"Updating existing conversation for session_id: {state.get('session_id')}")
            conversation_entry.messages = new_messages
            conversation_entry.last_updated = datetime.datetime.now()
        else:
            print(f"Creating new conversation for session_id: {state.get('session_id')}")
            new_conversation_entry = ConversationHistory(
                session_id=state.get("session_id"),
                messages=new_messages,
                last_updated=datetime.datetime.now()
            )
            db.add(new_conversation_entry)
        
        db.commit()
        print(f"Successfully saved conversation for session_id: {state.get('session_id')}")

    except Exception as e:
        db.rollback()
        print(f"***CRITICAL ERROR***: Failed to save data to database. Error: {e}")
    finally:
        db.close()
    
    return {
        "chat_history": new_messages,
        "telemetry_id": transaction_id
    }


# Build and compile the workflow
workflow = StateGraph(GraphState)
workflow.add_node("retrieve", retrieve_documents)
workflow.add_node("generate", generate_response)
workflow.set_entry_point("retrieve")
workflow.add_edge("retrieve", "generate")
workflow.add_edge("generate", END)
compiled_app = workflow.compile(checkpointer=memory)

# --- Lifespan Event Handler ---
@asynccontextmanager
async def lifespan(app: FastAPI):
    print("Application startup...")
    Base.metadata.create_all(bind=engine)
    yield
    print("Application shutdown...")

app.router.lifespan_context = lifespan

# --- 5. API Models ---
class ChatHistoryEntry(BaseModel):
    role: str
    content: str
    telemetry_id: Optional[str] = None

class ChatRequest(BaseModel):
    user_question: str
    session_id: str
    chat_history: Optional[List[ChatHistoryEntry]] = Field(default_factory=list)

class ChatResponse(BaseModel):
    ai_response: str
    updated_chat_history: List[ChatHistoryEntry]
    telemetry_entry_id: str

class FeedbackRequest(BaseModel):
    session_id: str
    telemetry_entry_id: str
    feedback_score: int
    feedback_text: Optional[str] = None

class ConversationSummary(BaseModel):
    session_id: str
    title: str

# --- 6. FastAPI Endpoints ---
@app.post("/upload-documents")
async def upload_documents(files: List[UploadFile] = File(...)):
    global vectorstore_retriever
    
    all_documents = []
    for uploaded_file in files:
        if uploaded_file.content_type == "text/plain":
            string_data = (await uploaded_file.read()).decode("utf-8")
            all_documents.append(Document(page_content=string_data, metadata={"source": uploaded_file.filename}))
        elif uploaded_file.content_type == "application/pdf":
            pdf_bytes = io.BytesIO(await uploaded_file.read())
            reader = PyPDF2.PdfReader(pdf_bytes)
            pdf_text = "".join([page.extract_text() + "\n" for page in reader.pages])
            all_documents.append(Document(page_content=pdf_text, metadata={"source": uploaded_file.filename}))
        elif uploaded_file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            docx_bytes = io.BytesIO(await uploaded_file.read())
            docx_docs = dx(docx_bytes)
            docx_content = "\n".join([para.text for para in docx_docs.paragraphs])
            all_documents.append(Document(page_content=docx_content, metadata={"source": uploaded_file.filename}))
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {uploaded_file.filename} ({uploaded_file.content_type})")
    
    if not all_documents:
        raise HTTPException(status_code=400, detail="No supported documents uploaded.")

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    splits = text_splitter.split_documents(all_documents)
    
    vectorstore = FAISS.from_documents(documents=splits, embedding=embeddings)
    vectorstore_retriever = vectorstore.as_retriever(search_kwargs={'k': 5})
    
    return {"message": f"Successfully processed {len(files)} documents and created knowledge base."}

@app.post("/chat", response_model=ChatResponse)
async def chat_with_rag(request: ChatRequest):
    global compiled_app
    global vectorstore_retriever
    if vectorstore_retriever is None:
        raise HTTPException(status_code=400, detail="Knowledge base not loaded. Please upload documents first.")

    initial_state = {
        "chat_history": [msg.model_dump() for msg in request.chat_history],
        "retrieved_documents": [],
        "user_question": request.user_question,
        "session_id": request.session_id
    }

    try:
        config = {"configurable": {"thread_id": request.session_id}}
        final_state = compiled_app.invoke(initial_state, config=config)
        
        ai_response_message = final_state["chat_history"][-1]["content"]
        updated_chat_history_dicts = final_state["chat_history"]

        return ChatResponse(
            ai_response=ai_response_message,
            updated_chat_history=updated_chat_history_dicts,
            telemetry_entry_id=final_state.get("telemetry_id")
        )
    except Exception as e:
        print(f"Internal Server Error: {e}")
        raise HTTPException(status_code=500, detail=f"An error occurred during chat processing: {e}")

@app.post("/feedback")
async def submit_feedback(request: FeedbackRequest):
    db = SessionLocal()
    try:
        telemetry_record = db.query(Telemetry).filter(
            Telemetry.transaction_id == request.telemetry_entry_id,
            Telemetry.session_id == request.session_id
        ).first()

        if not telemetry_record:
            raise HTTPException(status_code=404, detail="Telemetry entry not found or session ID mismatch.")

        existing_feedback = db.query(Feedback).filter(
            Feedback.telemetry_entry_id == request.telemetry_entry_id
        ).first()

        if existing_feedback:
            existing_feedback.feedback_score = request.feedback_score
            existing_feedback.feedback_text = request.feedback_text
            existing_feedback.timestamp = datetime.datetime.now()
        else:
            feedback_record = Feedback(
                telemetry_entry_id=request.telemetry_entry_id,
                feedback_score=request.feedback_score,
                feedback_text=request.feedback_text,
                user_query=telemetry_record.user_question,
                llm_response=telemetry_record.response,
                timestamp=datetime.datetime.now()
            )
            db.add(feedback_record)
            
        db.commit()

        return {"message": "Feedback submitted successfully."}

    except HTTPException as e:
        raise e
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"An error occurred: {str(e)}")
    finally:
        db.close()

@app.get("/conversations", response_model=List[ConversationSummary])
async def get_conversations():
    db = SessionLocal()
    try:
        conversations = db.query(ConversationHistory).order_by(ConversationHistory.last_updated.desc()).all()
        summaries = []
        for conv in conversations:
            first_user_message = next((msg for msg in conv.messages if msg["role"] == "user"), None)
            title = first_user_message["content"] if first_user_message else "New Conversation"
            summaries.append(ConversationSummary(session_id=conv.session_id, title=title[:30] + "..." if len(title) > 30 else title))
        return summaries
    finally:
        db.close()

@app.get("/conversations/{session_id}", response_model=List[ChatHistoryEntry])
async def get_conversation_history(session_id: str):
    db = SessionLocal()
    try:
        conversation = db.query(ConversationHistory).filter(ConversationHistory.session_id == session_id).first()
        if not conversation:
            raise HTTPException(status_code=404, detail="Conversation not found.")
        return conversation.messages
    finally:
        db.close()

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
